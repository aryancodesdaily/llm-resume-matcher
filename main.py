#-->imports

import os
import json
import time

from pathlib import Path
from dotenv import load_dotenv
from groq import Groq

from pydantic import BaseModel, Field\

from pypdf import PdfReader
from docx import Document


#-->basic code

load_dotenv()

my_api_key = os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("api key not found")

client = Groq(api_key=my_api_key)

model = "openai/gpt-oss-120b"

# paste the jobb description in "job_description"

job_description = """
Description

Do you want to solve real customer problems through innovative technology?
Do you enjoy working on scalable services in a collaborative team environment?
Do you want to see your code directly impact millions of customers worldwide?

At Amazon, we hire the best minds in technology to innovate and build on
behalf of our customers. Customer obsession is part of our company DNA,
which has made us one of the world's most beloved brands.

Our Software Development Engineers (SDEs) use modern technology to solve
complex problems while seeing their work's impact first-hand.

The challenges SDEs solve at Amazon are meaningful and influence millions
of customers, sellers, and products globally.

We seek individuals passionate about creating new products, features,
and services while managing ambiguity in an environment where development
cycles are measured in weeks, not years.

At Amazon, we believe in ownership at every level. As an SDE-I, you'll own
the entire lifecycle of your code - from design through deployment and
ongoing operations.

This ownership mindset, combined with our commitment to operational
excellence, ensures we deliver the highest quality solutions for our
customers.

We're looking for curious minds who think big and want to define tomorrow's
technology.

At Amazon, you'll grow into the high-impact engineer you know you can be,
supported by a culture of learning and mentorship.
"""



response_format={ "type":"json_object",}
# schema definitions

class Job(BaseModel):
    role: list[str] = Field(default_factory=list)
    required_skills: list[str] = Field(default_factory=list)
    preferred_skills: list[str] = Field(default_factory=list)
    minimum_experience: float | None = None
    education_requirement: list[str] = Field(default_factory=list)
    responsibility: list[str] = Field(default_factory=list)

    
class Result(BaseModel):
    score: float
    details: dict



class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = Field(default_factory=list)

    
class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None

    total_experience_years: float | None = None

    education: list[str] = Field(default_factory=list)
    skills: list[str] = Field(default_factory=list)
    projects: list[str] = Field(default_factory=list)
    experiences: list[Experience] = Field(default_factory=list)
    certifications: list[str] = Field(default_factory=list)



# function for parsing
def parse_job(job_description: str) -> Job:

    job_schema = Job.model_json_schema()

    system_prompt = f"""
    You are an expert HR assistant.

    Your task is to analyze a job description and extract structured
    information from it.

    Return ONLY valid JSON matching this schema:

    {json.dumps(job_schema, indent=2)}

    IMPORTANT:
    - Do not return the schema itself.
    - Do not return fields such as "properties", "title", or "type".
    - Extract actual information from the job description.
    - Do not invent information.
    - If minimum experience is not mentioned, return null.
    - If information for a list is missing, return an empty list.
    """


    user_prompt = f"""
    Analyze the following job description:

    {job_description}
    """

    messages = [
        {
            "role": "system",
            "content": system_prompt
        },
        {
            "role": "user",
            "content": user_prompt
        }
    ]
   
    response = client.chat.completions.create(
        model=model,
        messages=messages,
        response_format=response_format,
        temperature=0
    )

    
    raw_json = response.choices[0].message.content

    job_data = json.loads(raw_json)

    return Job(**job_data)


# functions for pdf reading

def read_pdf(file_path):

    reader = PdfReader(file_path)

    text = ""

    for page in reader.pages:

        page_text = page.extract_text()

        if page_text:
            text += page_text + "\n"

    return text

# function for docs file reading
def read_docx(file_path):

    document = Document(file_path)

    text = ""

    for paragraph in document.paragraphs:

        if paragraph.text.strip():
            text += paragraph.text + "\n"

    
    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                if cell.text.strip():
                    text += cell.text + "\n"

    return text



# function for read resume

def read_resume(file_path):

    extension = file_path.suffix.lower()

    if extension == ".pdf":
        return read_pdf(file_path)

    elif extension == ".docx":
        return read_docx(file_path)

    else:
        raise ValueError(
            f"Unsupported file type: {extension}"
        )



# after reading resume, function for parsing


def parse_resume(resume_text):

    resume_schema = Resume.model_json_schema()

    system_prompt = f"""
    You are an expert resume parser.

    Extract structured information from the candidate's resume.

    Return ONLY valid JSON matching this schema:

    {json.dumps(resume_schema, indent=2)}

    Rules:
    - Do not invent information.
    - Use null when a single-value field is not available.
    - Use an empty list when list information is not available.
    - Preserve the actual information from the resume.
    - Do not assume that a skill exists just because it is related to another skill.
    """

    user_prompt = f"""
    Parse the following resume:
    {resume_text}
    """

    messages = [
        {
            "role": "system",
            "content": system_prompt
        },
        {
            "role": "user",
            "content": user_prompt
        }
    ]

    response = client.chat.completions.create(
        model=model,
        messages=messages,
        response_format=response_format,
        temperature=0
    )

    raw_json = response.choices[0].message.content

    resume_data = json.loads(raw_json)

    return Resume(**resume_data)



def score(job, resume):

    result_schema = Result.model_json_schema()

    job_json = job.model_dump_json(indent=2)
    resume_json = resume.model_dump_json(indent=2)

    system_prompt = f"""
You are a strict technical recruiter and resume evaluator.

Compare the candidate's resume with the job requirements and estimate
the candidate's REALISTIC overall match.

The score must represent how well the candidate actually satisfies
the specific job requirements, not how impressive the resume looks.

IMPORTANT RULES:

1. Evaluate EVERY requirement individually.

2. For every required skill, determine whether it is:
   - Clearly demonstrated
   - Partially demonstrated
   - Not demonstrated

3. Do NOT give full credit merely because a skill appears in the
   resume's skills section. Look for actual evidence in projects,
   experience, or other relevant sections.

4. Professional experience is separate from projects.
   Projects, coursework, competitive programming, and college work
   must NOT be counted as professional experience.

5. If the job requires professional experience and the candidate
   does not have it, this should significantly reduce the overall
   score.

6. Evaluate project relevance based on:
   - relevance to the job
   - technical depth
   - number of projects
   - similarity to the responsibilities in the job

7. Preferred skills are weaker requirements and should have a
   smaller influence than required skills.

8. Do NOT assume that having many skills means a strong match.

9. Do NOT give a high score simply because the candidate has a
   strong general CS background.

10. Do NOT use arbitrary score ranges or fixed category weights.

11. The score must reflect the actual proportion and importance of
    requirements satisfied by the candidate.

12. A candidate missing major mandatory requirements should receive
    a clearly lower score even if many minor requirements are satisfied.

13. A candidate who satisfies only basic technical requirements but
    lacks required professional experience, relevant projects, or
    important mandatory skills should NOT receive a score above 70.

14. Do not invent information.

15. If a requirement says "A OR B", having either A or B satisfies
    that requirement.

After evaluating the requirements, provide:
- overall match score from 0 to 100
- matched requirements
- partially matched requirements
- missing important requirements
- experience assessment
- project assessment
- short final verdict

Return ONLY JSON matching this schema:

{json.dumps(result_schema, indent=2)}
"""




    user_prompt = f"""
    JOB REQUIREMENTS:
    {job_json}


    CANDIDATE RESUME:

    {resume_json}
    """

    messages = [
        {
            "role": "system",
            "content": system_prompt
        },
        {
            "role": "user",
            "content": user_prompt
        }
    ]

    response = client.chat.completions.create(
        model=model,
        messages=messages,
        response_format=response_format,
        temperature=0
    )

    raw_json = response.choices[0].message.content

    result_data = json.loads(raw_json)

    return Result(**result_data)



def main():

    print("Parsing job description...")

    job = parse_job(job_description)

    print("\nJob role:")
    print(job.role)

    print("\nMinimum experience:")
    print(job.minimum_experience)

    print("\nEducation:")
    print(job.education_requirement)

    print("Processing resumes...")
    resume_folder = Path("resumes")

    if not resume_folder.exists():
        raise FileNotFoundError(
            "The 'resumes' folder does not exist."
        )

    results = []

    for file_path in resume_folder.iterdir():

        if file_path.suffix.lower() not in [".pdf", ".docx"]:
            continue

        print(f"\nProcessing: {file_path.name}")

        try:

            
            resume_text = read_resume(file_path)

            if not resume_text.strip():
                print("Could not extract text. Skipping.")
                continue

            # Raw text → structured Resume
            parsed_resume = parse_resume(resume_text)

            time.sleep(2)

            # Resume + Job → Score
            result = score(job, parsed_resume)

            print("Score:", result.score)

            results.append(
                {
                    "name": parsed_resume.name,
                    "file": file_path.name,
                    "score": result.score,
                    "details": result.details
                }
            )

            time.sleep(2)

        except Exception as e:

            print(
                f"Error processing {file_path.name}: {e}"
            )

    # Sort highest score first
    results.sort(
        key=lambda x: x["score"],
        reverse=True
    )
    print("FINAL RANKING")
    for index, result in enumerate(results, start=1):

        print(
            f"\n{index}. "
            f"{result['name']} - "
            f"{result['score']}/100"
        )

        print(
            "Details:",
            result["details"]
        )

if __name__ == "__main__":
    main()